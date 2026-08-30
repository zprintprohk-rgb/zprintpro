#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Sales Contract Word 生成 (zh-hk 12 条款)
- 调用 python-docx (跨平台, OpenXML 标准)
- 12 条款按 §2.1 + §4 关键字段填充
- 双方信息 + 12 条款 + 签字盖章位 + 附件清单
- 落盘: F:\zprintpro-nextjs\docs\contracts\ZP-SA-2026-08-001.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

OUTPUT_DIR = r"F:\zprintpro-nextjs\docs\contracts"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# 字体配置 (zh-hk 用 標楷體 / 微軟正黑, en 用 Calibri)
ZH_FONT = "標楷體"
EN_FONT = "Calibri"

def set_cell_font(cell, text, font_name=ZH_FONT, size=10, bold=False, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # CJK 字体设置
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    return run

def set_para_font(paragraph, text, font_name=ZH_FONT, size=10, bold=False, color=None, align=None, line_spacing=1.5):
    paragraph.text = ""
    if align:
        paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    return run

def add_horizontal_line(paragraph):
    """添加水平线"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_border(cell, color="000000", size="6"):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

# ============================================================================
# 创建 Word 文档
# ============================================================================
doc = Document()

# 页面设置 (A4, 上下 2.54cm, 左右 3.17cm)
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# 默认字体 (Normal style)
style = doc.styles['Normal']
style.font.name = ZH_FONT
style.font.size = Pt(10)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.append(rFonts)
rFonts.set(qn('w:eastAsia'), ZH_FONT)
rFonts.set(qn('w:ascii'), EN_FONT)
rFonts.set(qn('w:hAnsi'), EN_FONT)

# ============================================================================
# 1. 封面 (Cover Page)
# ============================================================================
p_cover1 = doc.add_paragraph()
p_cover1.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_font(p_cover1, "", size=12)
p_cover1.paragraph_format.space_before = Pt(120)

p_cover_title = doc.add_paragraph()
p_cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_font(p_cover_title, "銷 售 合 同", size=28, bold=True, line_spacing=1.5)
p_cover_title.paragraph_format.space_before = Pt(0)
p_cover_title.paragraph_format.space_after = Pt(20)

p_cover_subtitle = doc.add_paragraph()
p_cover_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_font(p_cover_subtitle, "Sales Contract", size=16, bold=False, color=RGBColor(0x66, 0x66, 0x66), line_spacing=1.5)
p_cover_subtitle.paragraph_format.space_after = Pt(60)

# 合同基本信息表
table_info = doc.add_table(rows=4, cols=2)
table_info.alignment = WD_TABLE_ALIGNMENT.CENTER
table_info.columns[0].width = Cm(4)
table_info.columns[1].width = Cm(10)
info_data = [
    ("合同編號 / Contract No.", "ZP-SA-2026-08-001"),
    ("簽訂日期 / Signing Date", "2026 年 8 月 27 日"),
    ("簽訂地點 / Signing Place", "広東省深圳市龍崗區平湖街道嘉城路1號 (〒518111)"),
    ("適用法律 / Governing Law", "中華人民共和國民法典 + CISG + 香港特別行政區貨物售賣條例"),
]
for i, (k, v) in enumerate(info_data):
    set_cell_font(table_info.cell(i, 0), k, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_font(table_info.cell(i, 1), v, size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    for c in [table_info.cell(i, 0), table_info.cell(i, 1)]:
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(c)

# 封面底部 — 双方公司名
p_cover_blank = doc.add_paragraph()
set_para_font(p_cover_blank, "", size=12)
p_cover_blank.paragraph_format.space_before = Pt(80)

p_cover_seller = doc.add_paragraph()
p_cover_seller.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_font(p_cover_seller, "賣方: 深圳市彩龍印刷包裝有限公司", size=14, bold=True, line_spacing=2.0)

p_cover_arrow = doc.add_paragraph()
p_cover_arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_font(p_cover_arrow, "↓", size=14, bold=False, line_spacing=1.5)

p_cover_buyer = doc.add_paragraph()
p_cover_buyer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_font(p_cover_buyer, "買方: [客戶公司全稱]", size=14, bold=True, line_spacing=2.0)

# 分页
doc.add_page_break()

# ============================================================================
# 2. 正文 (Body) — 双方信息 + 12 条款
# ============================================================================
p_body_title = doc.add_paragraph()
p_body_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_font(p_body_title, "銷 售 合 同", size=18, bold=True, line_spacing=1.5)
p_body_title.paragraph_format.space_after = Pt(20)

# 序言
p_preamble = doc.add_paragraph()
set_para_font(
    p_preamble,
    "本合同由以下雙方於 2026 年 8 月 27 日在中華人民共和國廣東省深圳市簽訂。",
    size=11, line_spacing=1.75
)
p_preamble.paragraph_format.first_line_indent = Cm(0.74)
p_preamble.paragraph_format.space_after = Pt(10)

# --- 第一条 双方信息 ---
def add_clause_heading(num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    set_para_font(p, f"第 {num} 條　{title}", size=12, bold=True, line_spacing=1.5)

def add_clause_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.75
    p.paragraph_format.space_after = Pt(4)
    set_para_font(p, text, size=11, line_spacing=1.75)

add_clause_heading("一", "雙方信息")

# 卖方信息表
table_seller = doc.add_table(rows=7, cols=2)
table_seller.alignment = WD_TABLE_ALIGNMENT.CENTER
seller_data = [
    ("公司名稱", "深圳市彩龍印刷包裝有限公司"),
    ("英文名稱", "Shenzhen Cailong Printing Packaging Co., Ltd."),
    ("法定代表人", "唐運提 (Mr. Tang Yunti)"),
    ("註冊地址", "広東省深圳市龍崗區平湖街道嘉城路1號 (〒518111)"),
    ("聯繫電話", "+86 198 8085 1334 (WhatsApp 統一)"),
    ("電子郵箱", "zprintpro@outlook.com"),
    ("資質認證", "ISO 9001 質量管理體系 + FSC 森林認證 (Chain of Custody)"),
]
for i, (k, v) in enumerate(seller_data):
    set_cell_font(table_seller.cell(i, 0), k, size=10, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_font(table_seller.cell(i, 1), v, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    for c in [table_seller.cell(i, 0), table_seller.cell(i, 1)]:
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(c)
        c.paragraphs[0].paragraph_format.space_after = Pt(2)

p_blank1 = doc.add_paragraph()
set_para_font(p_blank1, "", size=10)

# 买方信息表
table_buyer = doc.add_table(rows=7, cols=2)
table_buyer.alignment = WD_TABLE_ALIGNMENT.CENTER
buyer_data = [
    ("公司名稱", "[客戶公司全稱]"),
    ("英文名稱", "[Customer Company Name]"),
    ("法定代表人", "[客戶代表人姓名]"),
    ("註冊地址", "[客戶地址]"),
    ("聯繫電話", "[客戶電話]"),
    ("電子郵箱", "[客戶郵箱]"),
    ("資質認證", "[客戶資質, 如適用]"),
]
for i, (k, v) in enumerate(buyer_data):
    set_cell_font(table_buyer.cell(i, 0), k, size=10, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_font(table_buyer.cell(i, 1), v, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    for c in [table_buyer.cell(i, 0), table_buyer.cell(i, 1)]:
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(c)
        c.paragraphs[0].paragraph_format.space_after = Pt(2)

# --- 第二条 标的 ---
add_clause_heading("二", "合同標的 (Subject Matter)")
add_clause_body(
    "2.1 品名: 客製化印刷品一批, 詳見附件 1《品名規格表》。"
)
add_clause_body(
    "2.2 規格: 按買方確認之樣品及設計文件執行, 符合 ISO 12647-2 CMYK 色彩標準 + FSC 認證紙材 + 客戶指定之工藝要求。"
)
add_clause_body(
    "2.3 數量: 詳見附件 1《品名規格表》 (預估 10,000 至 100,000 pcs 不等, 視品類)。"
)
add_clause_body(
    "2.4 單價: 詳見附件 1《品名規格表》 (含稅, 含基本包裝, 不含運輸及保險)。"
)
add_clause_body(
    "2.5 總價: 合計 HK$________ (港幣________整), 最終按實際驗收合格數量結算。"
)

# --- 第三条 质量标准 ---
add_clause_heading("三", "質量標準 (Quality Standard)")
add_clause_body(
    "3.1 賣方保證所交付印刷品符合以下標準: "
)
add_clause_body(
    "    (1) ISO 9001 質量管理體系;"
)
add_clause_body(
    "    (2) FSC 森林認證 (Chain of Custody), 紙材來源可追溯;"
)
add_clause_body(
    "    (3) CMYK 4 色印刷, 色彩符合 ISO 12647-2 標準 (FOGRA39 / GRACoL);"
)
add_clause_body(
    "    (4) 防水 / 防曬 / 耐磨 / 耐折 性能符合客戶要求 (按附件 2《技術規格》執行);"
)
add_clause_body(
    "    (5) FDA 21 CFR 176.170 (適用於食品接觸材料, 如客戶要求)。"
)
add_clause_body(
    "3.2 買方應在收到貨物 7 個工作天內完成驗收, 逾期未提出書面異議視為驗收合格。"
)

# --- 第四条 交期 ---
add_clause_heading("四", "交期 (Delivery Time)")
add_clause_body(
    "4.1 印刷生產週期: 3-5 個工作天 (急件 4-6 小時, 加急費 30-50%, 詳見附件 3《加急費標準》)。"
)
add_clause_body(
    "4.2 質檢 + 包裝: 1 個工作天。"
)
add_clause_body(
    "4.3 跨境物流: DHL 全球 2-4 個工作天 / FedEx 國際 3-5 個工作天 / SF International 5-7 個工作天 / 海運 15-30 天。"
)
add_clause_body(
    "4.4 總交期: 6-10 個工作天 (DHL, 含週末及法定節假日除外)。"
)
add_clause_body(
    "4.5 急件 SLA: 100 張起印 CMYK 全彩, 18:00 截單 → 翌日中午 12:00 前順豐送達香港, WhatsApp 30 秒即時報價。"
)

# --- 第五条 包装 ---
add_clause_heading("五", "包裝 (Packaging)")
add_clause_body(
    "5.1 內包裝: PE 袋 + 防潮珠 + 產品說明卡 (如需)。"
)
add_clause_body(
    "5.2 外包裝: 5 層 AB 浪瓦楞紙箱 + 順豐袋 + 緩衝氣泡袋。"
)
add_clause_body(
    "5.3 標識: 紙箱外標明買方名稱 + 合同編號 + 規格 + 數量 + 毛重 / 淨重 + 生產日期 + 有效期 (如適用)。"
)
add_clause_body(
    "5.4 特殊包裝: 禮盒 / 精裝盒 / 燙金盒 按附件 4《特殊包裝規格》執行。"
)

# --- 第六条 付款方式 ---
add_clause_heading("六", "付款方式 (Payment Terms)")
add_clause_body(
    "6.1 訂金: 合同簽訂後 3 個工作天內, 買方支付貨款總額 30% 作為訂金, 即 HK$________ (港幣________整)。"
)
add_clause_body(
    "6.2 尾款: 賣方發貨前 3 個工作天內, 買方支付貨款總額 70% 尾款, 即 HK$________ (港幣________整)。"
)
add_clause_body(
    "6.3 付款方式 (任選一): "
)
add_clause_body(
    "    (1) T/T 電匯 (優先, 推薦用於 USD / HKD / EUR / JPY 結算);"
)
add_clause_body(
    "    (2) L/C at sight 信用證 (適用於金額 ≥ USD 10,000 之大額交易);"
)
add_clause_body(
    "    (3) PayPal (適用於小金額 ≤ USD 5,000 之零售交易);"
)
add_clause_body(
    "    (4) Airwallex 跨境支付 (適用於港幣 / 美元 / 歐元結算);"
)
add_clause_body(
    "    (5) 銀行電匯 DBS Hong Kong (推薦, 支援港幣 / 美元 / 人民幣 / 歐元)。"
)
add_clause_body(
    "6.4 收款銀行信息: "
)
add_clause_body(
    "    收款人 (Beneficiary): 深圳市彩龍印刷包裝有限公司"
)
add_clause_body(
    "    收款銀行 (Bank): DBS Bank (Hong Kong) Limited"
)
add_clause_body(
    "    SWIFT Code: DHBKHKHH"
)
add_clause_body(
    "    銀行地址: 11/F, The Center, 99 Queen's Road Central, Hong Kong"
)
add_clause_body(
    "    帳戶號碼: [DBS HK USD / HKD / CNY 三帳戶, K3 必拍 1 次回復確認]"
)

# --- 第七条 验收 ---
add_clause_heading("七", "驗收 (Acceptance)")
add_clause_body(
    "7.1 買方應在收到貨物 7 個工作天內完成驗收, 並書面通知賣方驗收結果。"
)
add_clause_body(
    "7.2 驗收標準: 第三條質量標準 + 樣品確認書 + 附件 2《技術規格》。"
)
add_clause_body(
    "7.3 逾期未驗收: 買方在收到貨物 7 個工作天內未提出書面異議, 視為驗收合格, 賣方有權要求買方支付剩餘貨款。"
)
add_clause_body(
    "7.4 不合格處理: 經雙方確認不合格之印刷品, 賣方應在 5 個工作天內退貨退款或重新生產, 費用由賣方承擔。"
)

# --- 第八条 违约 ---
add_clause_heading("八", "違約責任 (Breach of Contract)")
add_clause_body(
    "8.1 賣方逾期交貨: 每日按逾期部分貨款 0.5% 支付違約金, 上限為 5% 貨款總額。"
)
add_clause_body(
    "8.2 賣方嚴重逾期: 逾期 ≥ 10 個工作天, 買方有權取消合同, 賣方退還已收貨款並支付 5% 違約金。"
)
add_clause_body(
    "8.3 買方逾期付款: 每日按逾期款項 0.5% 支付違約金, 上限為 5% 逾期款項。"
)
add_clause_body(
    "8.4 買方嚴重逾期: 逾期 ≥ 30 天, 賣方有權暫停生產 + 追究法律責任, 已收訂金不予退還。"
)
add_clause_body(
    "8.5 不可抗力: 因自然災害 / 戰爭 / 疫情 / 政府禁令等不可抗力, 雙方互不承擔違約責任, 但應及時通知對方並提供證明。"
)

# --- 第九条 知识产权 ---
add_clause_heading("九", "知識產權與保密 (IP & Confidentiality)")
add_clause_body(
    "9.1 設計版權: 買方提供之設計文件 / LOGO / 商標 / 品牌元素, 其知識產權歸買方所有, 賣方僅在合同範圍內使用。"
)
add_clause_body(
    "9.2 保密義務: 雙方對合同涉及之商業秘密 / 客戶資料 / 技術規格 / 價格信息 承擔保密義務, 保密期 2 年。"
)
add_clause_body(
    "9.3 賣方展示權: 賣方可在不透露買方機密信息前提下, 將本合同列為合作案例展示於官方網站 zprintpro.com (zh-hk / en / ja 8 locale)。"
)

# --- 第十条 争议解决 ---
add_clause_heading("十", "爭議解決 (Dispute Resolution)")
add_clause_body(
    "10.1 友好協商: 合同履行過程中發生爭議, 雙方應本著誠信原則友好協商解決。"
)
add_clause_body(
    "10.2 仲裁: 協商不成的, 任何一方均可提交中國國際經濟貿易仲裁委員會 (CIETAC) 深圳仲裁委員會按其現行有效仲裁規則進行仲裁。"
)
add_clause_body(
    "10.3 仲裁裁決為終局裁決, 對雙方均有約束力。"
)
add_clause_body(
    "10.4 仲裁費用: 敗訴方承擔。"
)

# --- 第十一条 适用法律 ---
add_clause_heading("十一", "適用法律 (Governing Law)")
add_clause_body(
    "11.1 本合同適用中華人民共和國法律 (含《民法典》《合同法》《對外貿易法》等)。"
)
add_clause_body(
    "11.2 國際貨物銷售適用《聯合國國際貨物銷售合同公約》(CISG)。"
)
add_clause_body(
    "11.3 香港特別行政區貨物售賣條例 (Cap. 26 Sale of Goods Ordinance) 作為補充適用法律。"
)

# --- 第十二条 其他 ---
add_clause_heading("十二", "其他條款 (Miscellaneous)")
add_clause_body(
    "12.1 合同生效: 本合同自雙方法定代表人或授權代表簽字並加蓋公章之日起生效。"
)
add_clause_body(
    "12.2 合同份數: 本合同一式兩份, 賣方與買方各執一份, 具有同等法律效力。"
)
add_clause_body(
    "12.3 附件: 本合同包含以下附件, 與本合同正文具有同等法律效力:"
)
add_clause_body(
    "    附件 1: 品名規格表 (Items Specification Table)"
)
add_clause_body(
    "    附件 2: 技術規格 (Technical Specification)"
)
add_clause_body(
    "    附件 3: 加急費標準 (Rush Fee Standard)"
)
add_clause_body(
    "    附件 4: 特殊包裝規格 (Special Packaging Specification)"
)
add_clause_body(
    "    附件 5: 樣品確認書 (Sample Confirmation)"
)
add_clause_body(
    "12.4 合同變更: 任何合同變更須經雙方書面同意。"
)
add_clause_body(
    "12.5 合同終止: 經雙方協商一致, 可書面終止本合同。"
)
add_clause_body(
    "12.6 未盡事宜: 本合同未盡事宜, 雙方另行協商並簽訂補充協議, 補充協議與本合同具有同等法律效力。"
)

# 分页
doc.add_page_break()

# ============================================================================
# 3. 签字页 (Signature Page)
# ============================================================================
p_sig_title = doc.add_paragraph()
p_sig_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_font(p_sig_title, "簽 字 蓋 章 頁", size=18, bold=True, line_spacing=1.5)
p_sig_title.paragraph_format.space_after = Pt(30)

p_sig_intro = doc.add_paragraph()
p_sig_intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_font(p_sig_intro, "Signature & Stamp Page", size=12, color=RGBColor(0x66, 0x66, 0x66), line_spacing=1.5)
p_sig_intro.paragraph_format.space_after = Pt(40)

# 双方签字盖章表
table_sig = doc.add_table(rows=7, cols=2)
table_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
table_sig.columns[0].width = Cm(7)
table_sig.columns[1].width = Cm(7)

# 表头
set_cell_font(table_sig.cell(0, 0), "賣方 (Seller)", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_font(table_sig.cell(0, 1), "買方 (Buyer)", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
for c in [table_sig.cell(0, 0), table_sig.cell(0, 1)]:
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(c)

# 主体
sig_data = [
    ("深圳市彩龍印刷包裝有限公司", "[客戶公司全稱]"),
    ("Shenzhen Cailong Printing Packaging Co., Ltd.", "[Customer Company Name]"),
    ("法定代表人: 唐運提 (Mr. Tang Yunti)", "法定代表人: [客戶代表人]"),
    ("", ""),
    ("簽字: ____________________", "簽字: ____________________"),
    ("日期: 2026 年 8 月 27 日", "日期: ____________________"),
]
for i, (left, right) in enumerate(sig_data):
    set_cell_font(table_sig.cell(i+1, 0), left, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_font(table_sig.cell(i+1, 1), right, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    for c in [table_sig.cell(i+1, 0), table_sig.cell(i+1, 1)]:
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(c)
        c.paragraphs[0].paragraph_format.space_after = Pt(6)

# 盖章位说明
p_stamp = doc.add_paragraph()
p_stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_font(p_stamp, "[公司公章]　　　　　　　　　　　　　[公司公章]", size=11, line_spacing=1.5)
p_stamp.paragraph_format.space_before = Pt(40)

# 页脚
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_font(p_footer, "— 本合同正本 —", size=10, color=RGBColor(0x99, 0x99, 0x99), line_spacing=1.5)
p_footer.paragraph_format.space_before = Pt(20)

# ============================================================================
# 页眉页脚
# ============================================================================
section = doc.sections[0]
header = section.header
p_header = header.paragraphs[0]
p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_para_font(p_header, "智印港 ZprintPro · 銷售合同 ZP-SA-2026-08-001", size=9, color=RGBColor(0x99, 0x99, 0x99))

footer = section.footer
p_footer = footer.paragraphs[0]
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_font(p_footer, "第 ", size=9, color=RGBColor(0x99, 0x99, 0x99))
# 页码字段
run = p_footer.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = "PAGE"
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run._element.append(fldChar1)
run._element.append(instrText)
run._element.append(fldChar2)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

run2 = p_footer.add_run(" 頁 / 共 ")
run2.font.size = Pt(9)
run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

run3 = p_footer.add_run()
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'begin')
instrText2 = OxmlElement('w:instrText')
instrText2.set(qn('xml:space'), 'preserve')
instrText2.text = "NUMPAGES"
fldChar4 = OxmlElement('w:fldChar')
fldChar4.set(qn('w:fldCharType'), 'end')
run3._element.append(fldChar3)
run3._element.append(instrText2)
run3._element.append(fldChar4)
run3.font.size = Pt(9)
run3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

run4 = p_footer.add_run(" 頁 · zprintpro.com")
run4.font.size = Pt(9)
run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ============================================================================
# 落盘
# ============================================================================
output_path = os.path.join(OUTPUT_DIR, "ZP-SA-2026-08-001.docx")
doc.save(output_path)
size_kb = os.path.getsize(output_path) / 1024
print(f"✅ Sales Contract Word 生成: {output_path}")
print(f"   Size: {size_kb:.1f} KB")
print(f"   Pages: 3 (封面 + 正文 12 條款 + 簽字蓋章頁)")
print(f"   Tables: 4 (合同信息表 + 雙方信息表 + 簽字蓋章表 + 賣方表 7 行 + 買方表 7 行)")
print(f"   Sections: 1 (with header + footer + page numbers)")
